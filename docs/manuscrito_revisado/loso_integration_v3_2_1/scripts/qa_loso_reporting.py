#!/usr/bin/env python3
"""CP7: numeric QA. Re-derives every cited figure from the canonical, frozen
loso_static_v1 outputs and cross-checks it against (a) loso_table6_source.csv,
(b) the actual Table 6 embedded in the clean manuscript, (c) the prose of the
new Results subsection 3.5. Prints PASS/FAIL per check; does not accept any
figure by assertion alone.
"""
from __future__ import annotations

from pathlib import Path

import docx
import pandas as pd

REPO_ROOT = Path(__file__).resolve().parents[4]
ANALYSIS_OUT = REPO_ROOT / "analysis" / "loso" / "outputs"
PHASE_DIR = Path(__file__).resolve().parents[1]
CLEAN_DOCX = REPO_ROOT / "docs" / "manuscrito_revisado" / "Manuscript_Methods_Results_English_Working_v9_10_LOSO_V3_2_1_clean.docx"

failures = []


def check(label, cond):
    status = "PASS" if cond else "FAIL"
    print(f"[{status}] {label}")
    if not cond:
        failures.append(label)


def main():
    summary = pd.read_csv(ANALYSIS_OUT / "loso_metrics_summary.csv")
    contrasts = pd.read_csv(ANALYSIS_OUT / "loso_contrasts.csv")
    run = pd.read_csv(ANALYSIS_OUT / "loso_metrics_by_run.csv")
    pred = pd.read_csv(ANALYSIS_OUT / "loso_predictions_long.csv")
    conv = pd.read_csv(ANALYSIS_OUT / "loso_convergence_summary.csv")

    check("condition summary rows == 16", len(summary) == 16)
    check("contrast rows == 12", len(contrasts) == 12)
    check("run rows == 48", len(run) == 48)
    check("prediction rows == 5580", len(pred) == 5580)
    check("convergence rows == 8", len(conv) == 8)

    auc_min, auc_max = summary["auc_point"].min(), summary["auc_point"].max()
    check("AUC min matches 44.4% (raw 0.4436..)", abs(auc_min - 0.443609022556391) < 1e-9)
    check("AUC max matches 63.9% (raw 0.6390..)", abs(auc_max - 0.6390374331550802) < 1e-9)

    incl = summary[(summary["auc_ci_low"] <= 0.5) & (summary["auc_ci_high"] >= 0.5)]
    excl = summary[~((summary["auc_ci_low"] <= 0.5) & (summary["auc_ci_high"] >= 0.5))]
    check("14/16 raw CI include 0.50", len(incl) == 14)
    excl_keys = set(zip(excl["held_out_site"], excl["roi_set"].astype(str), excl["model"]))
    check(
        "exceptions are exactly NYU-BNN12 and OHSU-BNN12",
        excl_keys == {("NYU", "12", "brainnetcnn"), ("OHSU", "12", "brainnetcnn")},
    )

    dim = contrasts[contrasts["contrast"] == "dimensionality"]
    check("4 dimensionality contrasts present", len(dim) == 4)
    check("all 4 dimensionality CI include zero", ((dim["delta_ci_low"] <= 0) & (dim["delta_ci_high"] >= 0)).all())
    dim_by_site = dim.set_index("held_out_site")["delta_point"]
    expected_dim_pp = {"NYU": -3.0, "Peking": -4.4, "NeuroIMAGE": -13.3, "OHSU": -13.5}
    for site, exp_pp in expected_dim_pp.items():
        got_pp = round(dim_by_site[site] * 100, 1)
        check(f"dimensionality delta {site} rounds to {exp_pp} pp (got {got_pp})", abs(got_pp - exp_pp) < 0.05)

    mf = contrasts[contrasts["contrast"].str.startswith("model_family")]
    check("8 model-family contrasts present", len(mf) == 8)
    check("all 8 model-family CI include zero", ((mf["delta_ci_low"] <= 0) & (mf["delta_ci_high"] >= 0)).all())
    check("12/12 total contrast CI include zero", ((contrasts["delta_ci_low"] <= 0) & (contrasts["delta_ci_high"] >= 0)).all())

    # --- Table 6 source CSV consistency ---
    t6 = pd.read_csv(PHASE_DIR / "loso_table6_source.csv")
    check("loso_table6_source.csv has 4 rows", len(t6) == 4)
    for _, r in t6.iterrows():
        site = r["held_out_site"]
        for prefix, roi, model in [
            ("bnn12", "12", "brainnetcnn"),
            ("logreg12", "12", "logreg"),
            ("bnn116", "116", "brainnetcnn"),
            ("logreg116", "116", "logreg"),
        ]:
            src = summary[
                (summary["held_out_site"] == site)
                & (summary["roi_set"].astype(str) == roi)
                & (summary["model"] == model)
            ].iloc[0]
            match = (
                abs(r[f"{prefix}_auc_point"] - src["auc_point"]) < 1e-12
                and abs(r[f"{prefix}_auc_ci_low"] - src["auc_ci_low"]) < 1e-12
                and abs(r[f"{prefix}_auc_ci_high"] - src["auc_ci_high"]) < 1e-12
            )
            check(f"table6_source {site}/{roi}/{model} matches canonical summary", match)

    # --- held-out N vs design.json (re-checked independently of build script) ---
    import json

    design = json.loads((REPO_ROOT / "results" / "loso" / "_design" / "loso_static_v1_design.json").read_text())
    expected_n = {"NYU": 177, "Peking": 183, "NeuroIMAGE": 39, "OHSU": 66}
    for site, n in expected_n.items():
        check(f"held-out n for {site} == {n}", design["rotation_sizes"][site]["test"] == n)
        check(f"table6_source held_out_n for {site} == {n}", int(t6.set_index("held_out_site").loc[site, "held_out_n"]) == n)

    # --- Table 6 as embedded in the clean docx vs loso_table6_source.csv ---
    d = docx.Document(CLEAN_DOCX)
    table6_docx = None
    for t in d.tables:
        header = [c.text for c in t.rows[0].cells]
        if header and header[0] == "Held-out site":
            table6_docx = t
            break
    check("Table 6 found embedded in clean docx", table6_docx is not None)
    if table6_docx is not None:
        check("embedded Table 6 has 5 rows (header + 4 sites)", len(table6_docx.rows) == 5)
        check("embedded Table 6 has 5 columns", len(table6_docx.columns) == 5)
        site_order = ["NYU", "Peking", "NeuroIMAGE", "OHSU"]
        for i, site in enumerate(site_order):
            row_cells = [c.text for c in table6_docx.rows[i + 1].cells]
            expected_label = f"{site} (held-out n={expected_n[site]})"
            check(f"embedded Table 6 row {i+1} label == '{expected_label}'", row_cells[0] == expected_label)
            src_row = t6.set_index("held_out_site").loc[site]
            expected_bnn12 = f"{src_row['bnn12_auc_point']*100:.1f} [{src_row['bnn12_auc_ci_low']*100:.1f}, {src_row['bnn12_auc_ci_high']*100:.1f}]"
            check(f"embedded Table 6 {site} BNN-12 cell == '{expected_bnn12}'", row_cells[1] == expected_bnn12)

    # --- Prose in section 3.5 vs canonical figures ---
    idx_35 = None
    for i, p in enumerate(d.paragraphs):
        if p.text.strip() == "Site-Held-Out Performance under Leave-One-Site-Out Evaluation":
            idx_35 = i
            break
    check("section 3.5 heading found in clean docx", idx_35 is not None)
    if idx_35 is not None:
        prose = " ".join(p.text for p in d.paragraphs[idx_35 : idx_35 + 6])
        check("prose states 44.4% and 63.9%", "44.4%" in prose and "63.9%" in prose)
        check("prose states 'Fourteen of the sixteen'", "Fourteen of the sixteen" in prose)
        for pp in ["-3.0 percentage points", "-4.4", "-13.3", "-13.5"]:
            check(f"prose contains dimensionality delta '{pp}'", pp in prose)
        check("prose states 'all eight'", "all eight" in prose)

        # Prohibited-claim scan (plan section 12.3): every hit must be classified,
        # not blanket-rejected -- a negated/allowed use (e.g. "No pooled estimate
        # ... was calculated", which D7/D11 require) is "allowed and scoped", not
        # an overclaim.
        allowed_context = {
            "pooled": "No pooled estimate across sites was calculated",
        }
        scan_terms = [
            "external validation", "independent validation", "generalizes", "generalizable",
            "future cohort", "scanner invariant", "site invariant", "equivalent", "noninferior",
            "non-inferior", "optimal", "best model", "superior", "clinical utility",
            "diagnostic tool", "biomarker", "pooled", "generalization gap", "identical architecture",
        ]
        unclassified_overclaims = []
        for term in scan_terms:
            if term in prose.lower():
                expected_ctx = allowed_context.get(term)
                if expected_ctx and expected_ctx.lower() in prose.lower():
                    print(f"  [scan] '{term}' found -> classified allowed_and_scoped (matches required negation: \"{expected_ctx}\")")
                else:
                    unclassified_overclaims.append(term)
        check("prose prohibited-claim scan: 0 unclassified overclaims", len(unclassified_overclaims) == 0)

    print()
    if failures:
        print(f"=== {len(failures)} CHECK(S) FAILED ===")
        for f in failures:
            print(" -", f)
        raise SystemExit(1)
    else:
        print("=== ALL CHECKS PASSED ===")


if __name__ == "__main__":
    main()
